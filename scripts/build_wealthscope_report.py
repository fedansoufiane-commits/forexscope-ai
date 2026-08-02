"""Build the five-page WealthScope AI 1.0 project report."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "Präsentation" / "Final"
DOCX_PATH = FINAL_DIR / "WealthScope_Ausarbeitung_5_Seiten_v1.0.docx"
DIAGNOSTICS_PATH = ROOT / "models" / "diagnostics.json"
EXPERIMENTS_PATH = ROOT / "models" / "validation_experiments.json"
REPO_URL = "https://github.com/fedansoufiane-commits/forexscope-ai"

GREEN = "1F5A4E"
GOLD = "B98519"
DARK = "17201D"
MUTED = "58635E"
LIGHT = "EAF1ED"
WHITE = "FFFFFF"


def de(value: float, digits: int = 3) -> str:
    """German decimal notation - the report is written in German throughout."""
    return f"{value:.{digits}f}".replace(".", ",")


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def add_text(doc: Document, text: str, *, bold=False, color=DARK, size=9.2,
             after=3, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.04
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc: Document, items: list[str], size=8.8) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.22)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(DARK)


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2 if level == 1 else 1)
    p.paragraph_format.space_after = Pt(4 if level == 1 else 2)
    keep_with_next(p)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(17 if level == 1 else 11)
    r.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else GOLD)


def add_page_label(doc: Document, kicker: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(kicker.upper())
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    add_heading(doc, title)


def add_info_box(doc: Document, title: str, text: str, color=LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_fill(cell, color)
    set_cell_margins(cell, 100, 130, 100, 130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    r2 = p2.add_run(text)
    r2.font.name = "Aptos"
    r2.font.size = Pt(8.6)
    r2.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_hyperlink(paragraph, url: str, text: str, size=8.6) -> None:
    """Clickable link - python-docx has no API for this, so build the XML."""
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    for tag, attrs in (("w:rFonts", {"w:ascii": "Aptos", "w:hAnsi": "Aptos"}),
                       ("w:b", {}),
                       ("w:color", {"w:val": GREEN}),
                       ("w:u", {"w:val": "single"}),
                       ("w:sz", {"w:val": str(int(size * 2))})):
        node = OxmlElement(tag)
        for key, value in attrs.items():
            node.set(qn(key), value)
        r_pr.append(node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


def add_access_box(doc: Document, title: str, url: str, lines: list[str],
                   color=LIGHT) -> None:
    """Info box whose first line is a real hyperlink."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_fill(cell, color)
    set_cell_margins(cell, 100, 130, 100, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    p_link = cell.add_paragraph()
    p_link.paragraph_format.space_after = Pt(2)
    p_link.paragraph_format.line_spacing = 1.0
    r_label = p_link.add_run("Quellcode und vollständige Anwendung:  ")
    r_label.font.name = "Aptos"
    r_label.font.size = Pt(8.6)
    r_label.font.color.rgb = RGBColor.from_string(DARK)
    add_hyperlink(p_link, url, url)

    for idx, line in enumerate(lines):
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_after = Pt(0 if idx == len(lines) - 1 else 1.5)
        p_line.paragraph_format.line_spacing = 1.0
        r_line = p_line.add_run(line)
        r_line.font.name = "Aptos"
        r_line.font.size = Pt(8.6)
        r_line.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, GREEN)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(header)
        r.bold = True
        r.font.name = "Aptos"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for row_no, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_margins(cells[idx])
            if row_no % 2:
                set_cell_fill(cells[idx], "F4F7F5")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.name = "Aptos"
            r.font.size = Pt(7.6)
            r.font.color.rgb = RGBColor.from_string(DARK)
            if widths:
                cells[idx].width = Cm(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("WEALTHSCOPE AI 1.0  ·  PROJEKTAUSARBEITUNG")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Soufiane Fedan  ·  IU Internationale Hochschule  ·  ")
    r.font.name = "Aptos"
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def build() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    diag = json.loads(DIAGNOSTICS_PATH.read_text(encoding="utf-8"))
    comparison = diag["model_comparison"]
    experiments = json.loads(EXPERIMENTS_PATH.read_text(encoding="utf-8"))
    splits = experiments["split_comparison"]
    split_summary = experiments["split_comparison_summary"]
    capacity_summary = experiments["capacity_sweep_summary"]

    doc = Document()
    configure_document(doc)

    # Seite 1
    add_text(doc, "WEALTHSCOPE AI", bold=True, color=GOLD, size=9, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Interaktive Finanzanalyse\nmit Machine Learning")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    add_text(
        doc,
        "Fünfseitige Projektausarbeitung · Version 1.0 · QUA³CK-Prozessmodell",
        bold=True, color=MUTED, size=10, after=8,
    )
    add_info_box(
        doc,
        "Leitfrage",
        "Wie können historische US-Aktienmarktdaten genutzt werden, um eine "
        "interaktive Finanzanalyse-App zu entwickeln, die technische Analyse, "
        "ML-Signalgebung und risikobasierte Positionsplanung nachvollziehbar kombiniert?",
    )
    add_heading(doc, "Kurzfassung", level=2)
    add_text(
        doc,
        "WealthScope AI ist ein wissenschaftlicher Streamlit-Prototyp. Er verbindet "
        "192.119 Feature-Datenpunkte aus 26 US-Blue-Chip-Titeln (1962–2017) mit "
        "technischen Indikatoren, einem reproduzierbaren Klassifikationsbenchmark "
        "und erklärbaren Risiko- und Transfermodulen. Ziel ist nicht die Behauptung "
        "sicherer Kursprognosen, sondern eine transparente Demonstration dessen, "
        "was Machine Learning aus historischen Kursdaten lernen kann – und wo seine "
        "Grenzen liegen.",
        size=9.1,
    )
    add_heading(doc, "Beitrag und Abgrenzung", level=2)
    add_bullets(doc, [
        "QUA³CK dokumentiert den Weg von der Frage bis zum Wissenstransfer.",
        "Fünf Modellgenerationen werden auf demselben zeitlichen Testfenster verglichen.",
        "StandardScaler und Imputation werden ausschließlich innerhalb der Trainingspipeline gefittet.",
        f"Das schwache Ergebnis wird nicht behauptet, sondern gegen die naheliegenden "
        f"Gegenerklärungen geprüft: Ein naiver Zufalls-Split würde dieselben Daten mit "
        f"AUC {de(split_summary['leaky_roc_auc'], 3)} statt "
        f"{de(split_summary['reference_roc_auc'], 3)} bewerten – ein "
        f"{de(split_summary['apparent_signal_factor'], 1)}-mal größeres scheinbares Signal.",
        "Die App erklärt Ergebnisse, bietet Export und Lernstudio, ist aber weder Handelsbot noch Anlageberatung.",
    ])
    add_text(
        doc,
        "Autor: Soufiane Fedan · Matrikelnummer 10247259 · Tutor: Klaus "
        "Quibeldey-Cirkel · Modul Data Analytics und Big Data",
        color=MUTED, size=8, after=0,
    )

    # Seite 2
    add_page_break(doc)
    add_page_label(doc, "Q · U · A", "Problem, Datenverständnis und Features")
    add_heading(doc, "1. Problemstellung und Hypothesen", level=2)
    add_text(
        doc,
        "Privatanleger treffen auf Datenfülle, widersprüchliche Signale und schwer "
        "erklärbare KI. WealthScope AI übersetzt Marktdaten in prüfbare Analysen. "
        "H1 untersucht, ob der Random Forest die Zufallsgrenze ROC-AUC 0,5 robust "
        "übertrifft. H2 prüft den verständlichen Wissenstransfer über Streamlit. "
        "H3 bewertet den praktischen Mehrwert der Kombination aus Technik, ML und Risiko.",
    )
    add_heading(doc, "2. Datenbasis und Datenqualität", level=2)
    add_table(doc, ["Aspekt", "Umsetzung", "Bedeutung"], [
        ["Quelle", "Kaggle US Stocks & ETFs, CC0", "Nachvollziehbare historische Basis"],
        ["Umfang", "192.119 Zeilen, 26 Ticker", "Mehrere Sektoren und Marktphasen"],
        ["Zeitraum", "1962–2017", "Lange Historie, aber Distribution Shift zu heute"],
        ["Ziel", "target_20d", "Kurs in 20 Handelstagen höher: 1, sonst 0"],
        ["Fehlwerte", "Median-Imputation in Pipeline", "Robust gegen Ausreißer; kein Testwissen"],
    ], widths=[2.8, 5.2, 8.2])
    add_heading(doc, "3. Feature Engineering", level=2)
    add_bullets(doc, [
        "Rendite: daily_return, return_5d und return_20d.",
        "Trend: relative Abstände zu MA-20, MA-50 und MA-200.",
        "Risiko: volatility_20d und drawdown zum bisherigen Hoch.",
        "Alle Features verwenden nur Informationen, die zum Beobachtungszeitpunkt vorliegen.",
    ])
    add_info_box(
        doc,
        "StandardScaler – zentraler methodischer Punkt",
        "z = (x − μTrain) / σTrain. Der Scaler lernt Mittelwert und Streuung nur aus "
        "X_train und transformiert X_test mit denselben Parametern. Das ist für "
        "Logistische Regression und Linear-SVM wichtig, weil deren Optimierung bzw. "
        "Abstände von der Merkmals-Skala abhängen. Entscheidungsbäume und Random "
        "Forests sind weitgehend skaleninvariant; dort ist der Scaler nicht erforderlich.",
        color="F7F1E3",
    )

    # Seite 3
    add_page_break(doc)
    add_page_label(doc, "A · A", "Algorithmus, Adaption und faire Validierung")
    add_heading(doc, "4. KI damals und heute: Modellvergleich", level=2)
    add_text(
        doc,
        "Der Benchmark zeigt die historische Entwicklung der Klassifikation: von "
        "einer einfachen Mehrheitsregel über lineare und regelbasierte Verfahren bis "
        "zum Ensemble. Alle Kandidaten erhalten dieselben Features und dieselben "
        "Trainings- und Testzeiträume.",
    )
    model_rows = []
    for key in ("dummy", "logistic", "decision_tree", "linear_svm", "random_forest"):
        item = comparison[key]
        scale = "ja" if key in {"logistic", "linear_svm"} else "nicht nötig"
        model_rows.append([
            str(item["year"]), item["label"], item["family"], scale,
            f"{de(item['test_metrics']['roc_auc'], 3)}",
        ])
    add_table(doc, ["Epoche", "Modell", "Grundidee", "Scaler", "AUC"], model_rows,
              widths=[1.7, 4.1, 4.6, 2.1, 1.4])
    add_text(
        doc,
        "Tabelle 1: Fünf Modellgenerationen auf identischem purged Out-of-Time-Fenster. "
        "Quelle: scripts/train_and_diagnose.py.",
        color=MUTED, size=7.5, after=4,
    )
    add_heading(doc, "5. Leakage-sichere Pipeline", level=2)
    add_info_box(
        doc,
        "Trainingsfolge",
        "Zeitlich sortieren → jüngste 20 % als unangetasteten Out-of-Time-Test "
        "reservieren → 20 Handelstage Purge vor jedem Validierungsfenster → "
        "Median-Imputer auf Training fitten → StandardScaler nur bei linearen "
        "Modellen auf Training fitten → Modell fitten → Test einmalig auswerten.",
    )
    add_heading(doc, "6. Warum kein zufälliger Split? Eine Messung", level=2)
    add_text(
        doc,
        "Die Zielvariable blickt 20 Handelstage voraus. Bei einem Random Split gelangen "
        "spätere Marktregime und überlappende Zielhorizonte in das Training. Wie groß "
        "dieser Effekt ist, bleibt in vielen Arbeiten unbeziffert – hier wurde er "
        "gemessen: dasselbe Modell, dieselben Daten, nur die Aufteilung variiert.",
    )
    add_table(doc, ["Aufteilung", "Leakage", "ROC-AUC", "Accuracy"], [
        [
            row["label"],
            row["leakage"],
            f"{de(row['roc_auc'], 4)}",
            f"{de(row['accuracy'], 4)} (Baseline {de(row['majority_baseline_accuracy'], 3)})",
        ]
        for row in splits
    ], widths=[5.0, 4.6, 2.2, 4.2])
    add_text(
        doc,
        f"Tabelle 2: Random Forest mit identischer Konfiguration. Gemessen am "
        f"Zufallsniveau 0,5 erscheint das Signal beim naiven Zufalls-Split mit "
        f"{de(split_summary['apparent_signal_leaky'], 3)} statt "
        f"{de(split_summary['apparent_signal_reference'], 3)} – also "
        f"{de(split_summary['apparent_signal_factor'], 1)}-mal größer, ohne dass sich "
        f"Modell oder Daten ändern. Quelle: scripts/validation_experiments.py.",
        color=MUTED, size=7.5, after=4,
    )
    add_text(
        doc,
        "Die niedrige Kennzahl dieser Arbeit ist damit kein Qualitätsmangel, sondern "
        "die Folge einer Validierung, die dem Modell die Zukunft entzieht. Vier "
        "expanding Walk-forward-Folds prüfen zusätzlich die Stabilität über frühere "
        "Zeitfenster.",
    )
    add_text(
        doc,
        "Random-Forest-Konfiguration: 200 Bäume, max_depth=8, "
        "min_samples_leaf=5, class_weight='balanced', random_state=42.",
        bold=True, color=GREEN, size=8.7,
    )

    # Seite 4
    add_page_break(doc)
    add_page_label(doc, "C · Evaluation", "Ergebnisse, Interpretation und Grenzen")
    add_table(doc, ["Modell", "ROC-AUC", "Balanced Accuracy", "Walk-forward AUC"], [
        [
            comparison[key]["label"],
            f"{de(comparison[key]['test_metrics']['roc_auc'], 3)}",
            f"{de(comparison[key]['test_metrics']['balanced_accuracy'], 3)}",
            f"{de(comparison[key]['walk_forward_roc_auc_mean'], 3)}",
        ]
        for key in ("dummy", "logistic", "decision_tree", "linear_svm", "random_forest")
    ], widths=[6.0, 3.2, 4.2, 3.5])
    add_text(
        doc,
        "Tabelle 3: Identisches Out-of-Time-Testfenster; Zufallsniveau der ROC-AUC = 0,5. "
        "Quelle: eigene Berechnung mit scripts/train_and_diagnose.py.",
        color=MUTED, size=7.5, after=4,
    )
    rf = comparison["random_forest"]
    add_table(doc, ["Kennzahl", "Random Forest", "Einordnung"], [
        ["Accuracy", f"{de(rf['test_metrics']['accuracy'], 3)}", "unter Mehrheitsbaseline 0,591"],
        ["Balanced Accuracy", f"{de(rf['test_metrics']['balanced_accuracy'], 3)}", "nur knapp über Zufall"],
        ["ROC-AUC", f"{de(rf['test_metrics']['roc_auc'], 3)}", "schwaches Ranking-Signal"],
        ["Walk-forward AUC", f"{de(rf['walk_forward_roc_auc_mean'], 3)}", "über Zeit nicht stabil stark"],
    ], widths=[4.0, 3.4, 8.0])
    add_heading(doc, "7. Bewertung der Hypothesen", level=2)
    add_bullets(doc, [
        "H1 falsifiziert, nicht ungeprüft: AUC 0,519 liegt über 0,5, der Abstand ist "
        "jedoch klein und zeitlich instabil. Die Hypothese war so formuliert, dass sie "
        "scheitern konnte – genau das macht sie prüfbar.",
        "H2 im Prototyp umgesetzt: Methodik, Kennzahlen, Export und Lernstudio sind zugänglich; eine formale Nutzerstudie steht aus.",
        "H3 teilweise plausibel: Die Kombination verbessert Orientierung und Risikokommunikation, beweist aber keinen wirtschaftlichen Mehrwert.",
    ])
    add_info_box(
        doc,
        "Zwei Gegenerklärungen wurden ausgeschlossen",
        f"Modellkapazität: Über sieben Regularisierungsstufen von unbeschränkt "
        f"(Trainings-AUC 1,000) bis stark gestutzt (0,543) bewegt sich der Test-AUC nur "
        f"um {de(capacity_summary['test_roc_auc_span'], 4)} "
        f"({de(capacity_summary['test_roc_auc_min'], 4)} bis "
        f"{de(capacity_summary['test_roc_auc_max'], 4)}); die beste Einstellung ist "
        f"{capacity_summary['best_label']}. Die Trainings-Test-Lücke entsteht also "
        f"ausschließlich auf der Trainingsseite – das Modell überanpasst Rauschen. "
        f"Datenmenge: Über die 2,3-fache Trainingsmenge verändert sich der "
        f"Validierungs-Score der Lernkurve nur um −0,005 und bleibt nahe "
        f"Zufallsniveau. Weder mehr Kapazität noch mehr Daten heben das Signal.",
        color="F7F1E3",
    )
    add_heading(doc, "8. Fachliche Grenzen", level=2)
    add_text(
        doc,
        "Historische Kurse enthalten wenig dauerhaftes Vorhersagesignal. Selektions- "
        "und Survivorship-Effekte, Transaktionskosten, Slippage, Steuern und "
        "Portfoliozwänge sind nicht modelliert. Live-Daten nach 2017 stellen einen "
        "Distribution Shift dar. Feature Importance und SHAP erklären statistische "
        "Zusammenhänge, aber keine Kausalität. Die ehrliche geringe Leistung ist "
        "deshalb das zentrale wissenschaftliche Ergebnis.",
    )

    # Seite 5
    add_page_break(doc)
    add_page_label(doc, "K · Knowledge", "Wissenstransfer, Kursbezug und Fazit")
    add_heading(doc, "9. Umsetzung in der Anwendung", level=2)
    add_table(doc, ["Baustein", "Funktion"], [
        ["Marktanalyse", "Kurs, Moving Averages, Candlesticks, Volatilität und Drawdown"],
        ["ML Insights", "Modellvergleich, Konfusionsmatrix, ROC/PR, Lernkurve, Wichtigkeiten"],
        ["Kapital-Kompass", "Risikobasierte Positionsgröße mit transparenten Annahmen"],
        ["Portfolio-Simulator", "Szenarien statt Renditeversprechen"],
        ["Lernstudio", "Quiz, direktes Feedback und arsnova.eu-kompatibler Export"],
        ["Methodik & Export", "QUA³CK, Model Card und reproduzierbare Dateien"],
    ], widths=[4.2, 11.8])
    add_access_box(
        doc,
        "Zugang zur Anwendung",
        REPO_URL,
        [
            "Lokal starten (Python 3.11 oder neuer):   1) git clone " + REPO_URL
            + ".git    2) cd forexscope-ai    3) pip install -r requirements.txt    "
            "4) streamlit run app.py",
            "Schritt 2 nicht überspringen: ohne den Wechsel ins Projektverzeichnis "
            "startet Schritt 4 eine fremde app.py oder bricht ab. Kontrolle: "
            "ls app.py src/ models/ muss alle drei auflisten.",
            "Die App öffnet sich anschließend unter http://localhost:8501. Die "
            "Diagnostik-Artefakte sind versioniert im Repository enthalten, ein "
            "Neutraining ist für den Start nicht erforderlich.",
            "Kennzahlen reproduzieren: python scripts/train_and_diagnose.py "
            "(Benchmark) und python scripts/validation_experiments.py "
            "(Tabelle 2 sowie der Kapazitäts-Sweep).",
        ],
    )
    add_heading(doc, "10. Bezug zur Vorlesung", level=2)
    add_text(
        doc,
        "Die Ausarbeitung verbindet ML-Grundlagen und historische Modellgenerationen "
        "mit Klassifikationsmetriken, Lernkurven, SVM, Entscheidungsbäumen und Random "
        "Forests. Der StandardScaler macht dabei sichtbar, dass Vorverarbeitung kein "
        "Nebenschritt ist: Sie beeinflusst lineare und abstandsbasierte Verfahren und "
        "muss innerhalb der Pipeline liegen. QUA³CK strukturiert die technische Arbeit "
        "und verhindert, dass das Modell vor Datenverständnis und Fragestellung steht.",
    )
    add_heading(doc, "11. Fazit", level=2)
    add_info_box(
        doc,
        "Kernaussage",
        "WealthScope AI 1.0 ist kein Kursorakel, sondern ein transparenter ML-Lern- "
        "und Analyseprototyp. Das Ergebnis der Arbeit ist nicht das Modell, sondern "
        "eine Messung: Wie viel verwertbare Information tragen rein kursbasierte "
        "technische Indikatoren? Antwort – nahezu keine, belegt an 190.527 "
        "Beobachtungen mit elf Jahren unangetastetem Testzeitraum und gegen zwei "
        "ausgeschlossene Gegenerklärungen. Das ist eine nachgerechnete, nicht bloß "
        "zitierte Bestätigung der Effizienzmarkthypothese. Der Mehrwert liegt in "
        "reproduzierbarer Methodik, fairem Modellvergleich und klarer "
        "Risikokommunikation. Der nächste wissenschaftliche Schritt ist ein "
        "Forward-Test mit Transaktionskosten, exogenen Variablen und formaler "
        "Nutzerstudie.",
    )
    add_heading(doc, "Quellen (Auswahl)", level=2)
    add_text(
        doc,
        "Fama, E. (1970): Efficient Capital Markets. Journal of Finance 25(2). · "
        "Géron, A. (2023): Praxiseinstieg Machine Learning, 3. Aufl. · "
        "Lundberg, S.; Lee, S.-I. (2017): A Unified Approach to Interpreting Model "
        "Predictions. NeurIPS. · Pedregosa, F. et al. (2011): Scikit-learn. JMLR 12. · "
        "Stock, A. et al. (2021): QUA³CK – A Machine Learning Development Process. · "
        "Marjanovic, B.: Huge Stock Market Dataset, Kaggle, CC0.",
        color=MUTED, size=7.7, after=0,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
